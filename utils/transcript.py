import os
import re
import logging
import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
from docx.shared import Pt
import asyncio

# Lazy load faster-whisper to save memory if not needed
def get_whisper_model():
    from faster_whisper import WhisperModel
    # Using 'base' model for a balance of speed and accuracy. 
    # 'int8' quantization is used to reduce RAM usage on CPU.
    return WhisperModel("base", device="cpu", compute_type="int8")

def get_video_id(url):
    """Extracts the video ID from a YouTube URL."""
    patterns = [
        r'(?:v=|\/|embed\/|youtu\.be\/)([0-9A-Za-z_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def fetch_transcript(url, temp_dir, progress_callback=None):
    """
    Fetches transcript for a YouTube video.
    First tries YouTubeTranscriptApi, then falls back to Faster-Whisper STT.
    """
    video_id = get_video_id(url)
    if not video_id:
        return None, "Video ID topilmadi."

    # 1. Try official captions/subtitles (Uzbek, Russian, English, Turkish)
    try:
        logging.info(f"Attempting to fetch official transcript for {video_id}")
        if progress_callback: progress_callback(5)
        
        # Fetch the list of available transcripts
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        
        # Try to find a manual or auto-generated transcript in our preferred languages
        # This will prefer manual over auto-generated if both exist.
        try:
            transcript = transcript_list.find_transcript(['uz', 'ru', 'en', 'tr'])
        except:
            # If not found, try to find any transcript and translate it to Uzbek (if possible) or just use whatever is available
            try:
                # Get the first available transcript
                available = list(transcript_list._manually_created_transcripts.values()) + \
                            list(transcript_list._generated_transcripts.values())
                if available:
                    transcript = available[0]
                else:
                    raise Exception("No transcripts available")
            except:
                raise Exception("No transcripts available")

        transcript_data = transcript.fetch()
        logging.info(f"Official transcript fetched ({transcript.language_code}).")
        return transcript_data, None
    except Exception as e:
        logging.info(f"Official transcript unavailable or failed: {e}. Falling back to AI Speech-to-Text.")

    # 2. Fallback: Download audio and use Faster-Whisper
    audio_path = os.path.join(temp_dir, f"{video_id}")
    
    # yt-dlp options for metadata and audio
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': audio_path,
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '128',
        }],
        'quiet': True,
        'no_warnings': True,
    }

    try:
        if progress_callback: progress_callback(10)
        logging.info(f"Checking video metadata for: {url}")
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # First, check video duration to prevent server overload
            info_dict = ydl.extract_info(url, download=False)
            duration = info_dict.get('duration', 0)
            
            # Limit: 1.5 hours (90 minutes) for AI transcription
            if duration > 5400: 
                return None, "Video juda uzun. AI transkripsiya uchun maksimal davomiylik 1.5 soat (90 daqiqa) qilib belgilangan."

            logging.info(f"Downloading audio for STT fallback (Duration: {duration}s)")
            ydl.download([url])
        
        final_audio_path = audio_path + ".mp3"
        if not os.path.exists(final_audio_path):
            # Check if it was downloaded with a different extension or directly
            if os.path.exists(audio_path):
                final_audio_path = audio_path
            else:
                return None, "Audioni yuklab olishda xatolik yuz berdi."

        if progress_callback: progress_callback(30)
        logging.info("Starting AI transcription using Faster-Whisper...")
        
        # Load model and transcribe
        model = get_whisper_model()
        # transcribe() segments is a generator
        segments, info = model.transcribe(final_audio_path, beam_size=5)
        
        transcript_data = []
        for segment in segments:
            transcript_data.append({
                'text': segment.text.strip(),
                'start': segment.start,
                'duration': segment.end - segment.start
            })
            if info.duration > 0 and progress_callback:
                pct = 30 + (segment.end / info.duration * 65)
                progress_callback(min(95, pct))

        logging.info("AI transcription completed.")
        return transcript_data, None

    except Exception as e:
        logging.error(f"STT Fallback failed: {e}")
        return None, f"Xatolik: {str(e)}"
    finally:
        # Cleanup audio file after transcription
        try:
            if 'final_audio_path' in locals() and os.path.exists(final_audio_path):
                os.remove(final_audio_path)
        except:
            pass

def create_transcript_docx(transcript_data, output_path):
    """Formats transcript list into a professional Word document."""
    try:
        doc = Document()
        
        # Set Document Style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        doc.add_heading('Video Transkripti', 0)
        
        if not transcript_data:
            doc.add_paragraph("Transkript mavjud emas.")
        else:
            # Join all text segments into one or more paragraphs without timestamps
            # For better readability, we'll combine them but keep some structure
            current_paragraph = ""
            for i, entry in enumerate(transcript_data):
                current_paragraph += entry['text'] + " "
                
                # Every 5 segments or at the end, start a new paragraph to avoid one giant block
                if (i + 1) % 5 == 0 or i == len(transcript_data) - 1:
                    doc.add_paragraph(current_paragraph.strip())
                    current_paragraph = ""
        
        doc.save(output_path)
        return True
    except Exception as e:
        logging.error(f"Error creating docx: {e}")
        return False
