import os
import logging
import traceback
import sys
import subprocess
import shutil

def word_to_pdf(docx_path: str, output_pdf_path: str, progress_callback=None) -> bool:
    """
    Converts a Word (.docx) file to a PDF file.
    On Windows it uses docx2pdf (requires MS Word).
    On Linux it uses LibreOffice CLI (requires libreoffice to be installed).
    """
    try:
        logging.info(f"Starting Word to PDF conversion: {docx_path}")
        if progress_callback: progress_callback(20)
        
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(output_pdf_path)
        outdir = os.path.dirname(abs_pdf)
        
        if sys.platform == "win32":
            from docx2pdf import convert
            convert(abs_docx, abs_pdf)
            return os.path.exists(abs_pdf)
        else:
            # Use a temporary user profile for LibreOffice to avoid profile locks/hangs in service mode
            # IMPORTANT: Path must be absolute for -env:UserInstallation
            profile_dir = os.path.abspath(os.path.join(outdir, f"lo_profile_{os.getpid()}"))
            os.makedirs(profile_dir, exist_ok=True)
            
            cmd = [
                'libreoffice', 
                '--headless', 
                '--nologo',
                '--nodefault',
                '--norestore',
                '--nolockcheck',
                f'-env:UserInstallation=file://{profile_dir}', 
                '--convert-to', 'pdf', 
                abs_docx, 
                '--outdir', outdir
            ]
            
            logging.info(f"Running LibreOffice with Profile: {profile_dir}")
            if progress_callback: progress_callback(40)
            
            try:
                # Add a timeout to prevent hanging the whole bot queue
                result = subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
                logging.info("LibreOffice command completed successfully")
            except subprocess.TimeoutExpired:
                logging.error(f"LibreOffice conversion timed out after 60 seconds: {abs_docx}")
                return False
            except subprocess.CalledProcessError as e:
                error_msg = e.stderr.decode('utf-8', errors='replace')
                logging.error(f"LibreOffice failed with error: {error_msg}")
                return False
            finally:
                # Clean up temp profile
                try: shutil.rmtree(profile_dir)
                except Exception as e: 
                    logging.warning(f"Could not clean up LO profile {profile_dir}: {e}")
            
            if progress_callback: progress_callback(80)
            
            base_name = os.path.splitext(os.path.basename(abs_docx))[0]
            generated_pdf = os.path.join(outdir, f"{base_name}.pdf")
            
            if generated_pdf != abs_pdf and os.path.exists(generated_pdf):
                os.rename(generated_pdf, abs_pdf)
            
            if progress_callback: progress_callback(100)
            return os.path.exists(abs_pdf)
    except Exception as e:
        logging.error(f"Error converting Word to PDF: {e}")
        return False

def pdf_to_word(pdf_path: str, output_docx_path: str, progress_callback=None) -> bool:
    """
    Converts a PDF file to a Word (.docx) file with table support and OCR fallback.
    Maintains layout by tracking table positions relative to text.
    """
    try:
        logging.info(f"Starting Enhanced PDF to Word conversion: {pdf_path}")
        if progress_callback: progress_callback(5)
        
        import fitz
        import pdfplumber
        import pytesseract
        from PIL import Image
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set default font for multilingual support (Uzbek, Russian, English, Turkish)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Open with both fitz (for pixmaps) and pdfplumber (for tables/text)
        pdf_fitz = fitz.open(pdf_path)
        pdf_plumber = pdfplumber.open(pdf_path)
        
        total_pages = len(pdf_plumber.pages)
        logging.info(f"PDF loaded: {total_pages} pages found")
        
        for i, page in enumerate(pdf_plumber.pages):
            logging.info(f"Processing page {i+1}/{total_pages}")
            
            # Get pixmap for OCR fallback
            pix = pdf_fitz[i].get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Scaling factor between pdfplumber coordinates (pts) and pixmap pixels
            scale_x = pix.width / page.width
            scale_y = pix.height / page.height
            
            # 1. Find tables
            tables = page.find_tables()
            table_bboxes = [t.bbox for t in tables]
            
            # 2. Get text objects not in tables
            words = page.extract_words()
            non_table_words = []
            for w in words:
                in_table = False
                for bbox in table_bboxes:
                    # w['x0'], w['top'], w['x1'], w['bottom']
                    if w['x0'] >= bbox[0] and w['top'] >= bbox[1] and w['x1'] <= bbox[2] and w['bottom'] <= bbox[3]:
                        in_table = True
                        break
                if not in_table:
                    non_table_words.append(w)
            
            # Group non-table words into paragraphs
            page_content = []
            if non_table_words:
                # Simple grouping by vertical proximity
                non_table_words.sort(key=lambda x: (x['top'], x['x0']))
                current_para = []
                last_bottom = -1
                tolerance = 10 # pts
                
                for w in non_table_words:
                    if last_bottom == -1 or w['top'] <= last_bottom + tolerance:
                        current_para.append(w)
                    else:
                        if current_para:
                            current_para.sort(key=lambda x: x['x0'])
                            text = " ".join([word['text'] for word in current_para])
                            page_content.append({'type': 'text', 'content': text, 'top': current_para[0]['top']})
                        current_para = [w]
                    last_bottom = w['bottom']
                
                if current_para:
                    current_para.sort(key=lambda x: x['x0'])
                    text = " ".join([word['text'] for word in current_para])
                    page_content.append({'type': 'text', 'content': text, 'top': current_para[0]['top']})
            
            # If no words found (scanned PDF), OCR the whole page but mask tables if any
            if not words and not tables:
                text = pytesseract.image_to_string(img, lang='uzb+rus+eng+tur')
                if text.strip():
                    page_content.append({'type': 'text', 'content': text.strip(), 'top': 0})
            elif not words and tables:
                # Scanned PDF with vector tables? Or just no text found.
                # We should still OCR the non-table areas.
                # For simplicity, we'll OCR the whole page if no words at all.
                text = pytesseract.image_to_string(img, lang='uzb+rus+eng+tur')
                if text.strip():
                    # This might overlap with tables, but it's a fallback
                    page_content.append({'type': 'text', 'content': text.strip(), 'top': 0})

            # Add tables to page content
            for t in tables:
                page_content.append({'type': 'table', 'obj': t, 'top': t.bbox[1]})
            
            # Sort everything by top coordinate to maintain layout
            page_content.sort(key=lambda x: x['top'])
            
            # 3. Add to docx
            for item in page_content:
                if item['type'] == 'text':
                    if item['content'].strip():
                        p = doc.add_paragraph(item['content'])
                        # Ensure font is applied to the run
                        for run in p.runs:
                            run.font.name = 'Arial'
                elif item['type'] == 'table':
                    t_obj = item['obj']
                    table_rows = t_obj.rows
                    if not table_rows: continue
                    
                    num_rows = len(table_rows)
                    num_cols = len(table_rows[0].cells) if num_rows > 0 else 0
                    
                    if num_rows > 0 and num_cols > 0:
                        docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                        docx_table.style = 'Table Grid'
                        
                        # Set column widths
                        for c_idx in range(num_cols):
                            cell_bbox = table_rows[0].cells[c_idx]
                            if cell_bbox:
                                width_pts = cell_bbox[2] - cell_bbox[0]
                                docx_table.columns[c_idx].width = Inches(width_pts / 72.0)
                        
                        for r_idx, row_obj in enumerate(table_rows):
                            for c_idx, cell_bbox in enumerate(row_obj.cells):
                                if cell_bbox is None: continue
                                
                                cell = docx_table.cell(r_idx, c_idx)
                                
                                # Extract text from pdfplumber
                                cell_area = page.within_bbox(cell_bbox)
                                cell_text = cell_area.extract_text()
                                
                                # OCR Fallback for scanned/image cells
                                if not cell_text or not cell_text.strip():
                                    crop_bbox = (
                                        cell_bbox[0] * scale_x,
                                        cell_bbox[1] * scale_y,
                                        cell_bbox[2] * scale_x,
                                        cell_bbox[3] * scale_y
                                    )
                                    # Ensure crop_bbox is within image bounds
                                    crop_bbox = (
                                        max(0, crop_bbox[0]),
                                        max(0, crop_bbox[1]),
                                        min(pix.width, crop_bbox[2]),
                                        min(pix.height, crop_bbox[3])
                                    )
                                    if crop_bbox[2] > crop_bbox[0] and crop_bbox[3] > crop_bbox[1]:
                                        cell_img = img.crop(crop_bbox)
                                        # Use multiple languages for OCR
                                        cell_text = pytesseract.image_to_string(cell_img, lang='uzb+rus+eng+tur').strip()
                                
                                cell.text = (cell_text or "").strip()
                                # Ensure font is applied to cell content
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
            
            if i < total_pages - 1:
                doc.add_page_break()
            
            if progress_callback:
                progress_callback(10 + ((i + 1) / total_pages * 85))

        doc.save(output_docx_path)
        pdf_fitz.close()
        pdf_plumber.close()
        if progress_callback: progress_callback(100)
        return os.path.exists(output_docx_path)
    except Exception as e:
        logging.error(f"Error in enhanced pdf_to_word: {e}\n{traceback.format_exc()}")
        return False

def md_to_pdf(md_path: str, output_pdf_path: str, progress_callback=None) -> bool:
    """
    Converts a Markdown (.md) file to a PDF file.
    """
    try:
        logging.info(f"Starting Markdown to PDF conversion: {md_path}")
        if progress_callback: progress_callback(10)
        import markdown
        from xhtml2pdf import pisa
        import os
        
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
            
        import base64, zlib, re
        def replacer(match):
            mermaid_code = match.group(1).strip()
            compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
            b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
            return f"![Mermaid Diagram](https://kroki.io/mermaid/png/{b64})"
            
        md_content = re.sub(r'```mermaid\s*\n(.*?)\n```', replacer, md_content, flags=re.DOTALL | re.IGNORECASE)
        if progress_callback: progress_callback(40)
            
        html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite', 'tables'])
        if progress_callback: progress_callback(60)
        
        script_dir = os.path.dirname(os.path.abspath(__file__))
        font_dir = os.path.join(script_dir, "fonts")
        font_path = os.path.join(font_dir, "Roboto-Regular.ttf").replace('\\', '/')
        
        if not os.path.exists(font_path):
            import urllib.request
            os.makedirs(font_dir, exist_ok=True)
            font_url = "https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf"
            urllib.request.urlretrieve(font_url, font_path)
        
        css = f"""
        <style>
            @font-face {{
                font-family: 'Roboto';
                src: url('{font_path}');
            }}
            @page {{ margin: 2cm; }}
            body {{ font-family: 'Roboto', Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333; }}
            h1, h2, h3, h4, h5, h6 {{ font-family: 'Roboto'; color: #222; margin-top: 1.5em; margin-bottom: 0.5em; }}
            code {{ font-family: 'Roboto', monospace; background-color: #f8f9fa; padding: 2px 4px; border-radius: 4px; font-size: 0.9em; }}
            pre {{ background-color: #f8f9fa; padding: 12px; border-radius: 4px; white-space: pre-wrap; font-family: 'Roboto', monospace; font-size: 0.9em; border: 1px solid #e9ecef; }}
            blockquote {{ font-family: 'Roboto'; border-left: 4px solid #adb5bd; padding-left: 15px; color: #6c757d; font-style: italic; margin-left: 0; }}
            table {{ font-family: 'Roboto'; border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
            th, td {{ font-family: 'Roboto'; border: 1px solid #dee2e6; padding: 8px; text-align: left; }}
            th {{ background-color: #e9ecef; font-weight: bold; }}
            a {{ color: #0d6efd; text-decoration: none; }}
            img {{ max-width: 100%; height: auto; }}
            hr {{ border: 0; border-top: 1px solid #eee; margin: 20px 0; }}
        </style>
        """
        full_html = f"<html><head>{css}</head><body>{html_content}</body></html>"
        
        with open(output_pdf_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(full_html, dest=result_file)
        
        if progress_callback: progress_callback(100)
        return not pisa_status.err
    except Exception as e:
        logging.error(f"Error in md_to_pdf: {e}\n{traceback.format_exc()}")
        return False
