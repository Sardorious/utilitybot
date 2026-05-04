import sys
import os
import logging
import fitz

# Setup logging
logging.basicConfig(level=logging.INFO)

# Add current dir to path
sys.path.append(os.getcwd())

from utils.converter import pdf_to_word

def create_test_pdf(filename):
    print(f"Creating test PDF: {filename}")
    doc = fitz.open()
    page = doc.new_page()
    
    # Text before
    page.insert_text((50, 50), "Text Before Table", fontsize=12)
    
    # Table layout (bbox: x0, y0, x1, y1)
    # Header
    page.draw_rect((50, 70, 450, 100), color=(0,0,0), width=1)
    page.draw_line((250, 70), (250, 100), color=(0,0,0), width=1)
    page.insert_text((60, 90), "Header A", fontsize=11)
    page.insert_text((260, 90), "Header B", fontsize=11)
    
    # Row 1
    page.draw_rect((50, 100, 450, 130), color=(0,0,0), width=1)
    page.draw_line((250, 100), (250, 130), color=(0,0,0), width=1)
    page.insert_text((60, 120), "Value A1", fontsize=10)
    page.insert_text((260, 120), "Value B1", fontsize=10)
    
    # Text after
    page.insert_text((50, 160), "Text After Table", fontsize=12)
    
    doc.save(filename)
    doc.close()
    print("PDF created successfully.")

def run_test():
    pdf_path = "manual_test.pdf"
    docx_path = "manual_test.docx"

    create_test_pdf(pdf_path)

    print("\n--- Starting PDF to Word ---")
    def progress(p):
        print(f"Progress: {p:.1f}%")

    if pdf_to_word(pdf_path, docx_path, progress_callback=progress):
        print(f"Successfully created {docx_path}")
        
        from docx import Document
        doc = Document(docx_path)
        print(f"\nResults:")
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
        
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                print(f"P{i}: {p.text}")
        
        for i, t in enumerate(doc.tables):
            print(f"\nTable {i}:")
            for r in t.rows:
                print(f"| {' | '.join([c.text for c in r.cells])} |")
    else:
        print("Conversion failed.")

if __name__ == "__main__":
    run_test()
